# -*- coding: utf-8 -*-
"""
Genera la Especificación de Requisitos de Software (ERS / SRS) del
Sistema de Gestión de Inventario JJ Congeminco, en formato Word (.docx).

Documento de ingeniería inversa: describe los requisitos funcionales y no
funcionales tal como se habrían definido antes de construir el sistema,
derivados del análisis del código existente.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- Paleta de marca ----------
NARANJA = RGBColor(0xE8, 0x63, 0x1A)
OSCURO = RGBColor(0x1E, 0x25, 0x30)
GRIS = RGBColor(0x5B, 0x65, 0x73)
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)
VERDE = RGBColor(0x13, 0x7A, 0x4B)

doc = Document()

# ---------- Estilos base ----------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)

def set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    trPr.append(th)

def h(text, level=1, color=NARANJA):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    run.font.color.rgb = color
    run.font.name = "Calibri"
    if level == 1:
        run.font.size = Pt(17)
    elif level == 2:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)
    return p

def para(text="", bold=False, italic=False, size=10.5, color=None, align=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def numbered(text):
    return doc.add_paragraph(text, style="List Number")

def kv_table(rows, widths=(4.0, 12.0)):
    """Tabla de 2 columnas clave/valor."""
    t = doc.add_table(rows=0, cols=2)
    t.style = "Light Grid Accent 2"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for k, v in rows:
        cells = t.add_row().cells
        cells[0].text = ""
        rp = cells[0].paragraphs[0].add_run(k)
        rp.bold = True
        cells[1].text = str(v)
    for row in t.rows:
        row.cells[0].width = Cm(widths[0])
        row.cells[1].width = Cm(widths[1])
    doc.add_paragraph()
    return t

def grid_table(headers, data, col_widths=None, header_bg="1E2530"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, txt in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(txt)
        run.bold = True
        run.font.color.rgb = BLANCO
        run.font.size = Pt(9.5)
        set_cell_bg(hdr[i], header_bg)
    set_repeat_header(t.rows[0])
    for fila in data:
        cells = t.add_row().cells
        for i, val in enumerate(fila):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t

def rf_block(rid, nombre, prioridad, descripcion, actores, precond, reglas):
    """Ficha de requisito funcional."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run(f"{rid} — {nombre}")
    r.bold = True
    r.font.size = Pt(11.5)
    r.font.color.rgb = OSCURO
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    def add(k, v):
        cells = t.add_row().cells
        cells[0].text = ""
        rr = cells[0].paragraphs[0].add_run(k)
        rr.bold = True
        rr.font.size = Pt(9.5)
        set_cell_bg(cells[0], "FDEEE3")
        cells[1].text = ""
        if isinstance(v, list):
            for j, item in enumerate(v):
                pp = cells[1].paragraphs[0] if j == 0 else cells[1].add_paragraph()
                run = pp.add_run("• " + item)
                run.font.size = Pt(9.5)
        else:
            run = cells[1].paragraphs[0].add_run(str(v))
            run.font.size = Pt(9.5)
        cells[0].width = Cm(3.2)
        cells[1].width = Cm(12.8)
    add("Prioridad", prioridad)
    add("Descripción", descripcion)
    add("Actores", actores)
    add("Precondiciones", precond)
    add("Reglas y validaciones", reglas)
    doc.add_paragraph()

def page_break():
    doc.add_page_break()

# ============================================================
# PORTADA
# ============================================================
for _ in range(3):
    doc.add_paragraph()

p = para("ESPECIFICACIÓN DE REQUISITOS DE SOFTWARE", bold=True, size=24, color=OSCURO,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para("(ERS / SRS)", bold=True, size=14, color=NARANJA, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

para("Sistema de Gestión de Inventario", bold=True, size=20, color=OSCURO,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("JJ Congeminco — Líderes en Soluciones", size=13, color=GRIS,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)

# Recuadro de metadatos
meta = doc.add_table(rows=0, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
meta.style = "Light List Accent 2"
for k, v in [
    ("Proyecto", "Sistema de inventario JJ Congeminco (web)"),
    ("Plataforma", "inventario.congeminco.com"),
    ("Tipo de documento", "Especificación de Requisitos (ingeniería inversa)"),
    ("Estándar de referencia", "Inspirado en IEEE 830 / ISO-IEC-IEEE 29148"),
    ("Versión del documento", "1.0"),
    ("Fecha", "Junio 2026"),
    ("Estado", "Línea base (baseline)"),
    ("Confidencialidad", "Uso interno"),
]:
    cells = meta.add_row().cells
    cells[0].text = ""
    rr = cells[0].paragraphs[0].add_run(k)
    rr.bold = True
    rr.font.size = Pt(10)
    cells[1].text = v
    cells[1].paragraphs[0].runs[0].font.size = Pt(10)
    cells[0].width = Cm(5)
    cells[1].width = Cm(11)

for _ in range(6):
    doc.add_paragraph()
para("Documento generado a partir del análisis del sistema en producción, "
     "reconstruyendo los requisitos funcionales y no funcionales como si se "
     "hubieran definido al inicio del proyecto.", italic=True, size=9.5, color=GRIS,
     align=WD_ALIGN_PARAGRAPH.CENTER)
page_break()

# ============================================================
# HISTORIAL DE CAMBIOS
# ============================================================
h("Historial de versiones", level=1)
grid_table(
    ["Versión", "Fecha", "Autor", "Descripción"],
    [
        ["1.0", "Jun 2026", "Equipo de Proyecto", "Línea base. Reconstrucción de requisitos a partir del sistema implementado."],
    ],
    col_widths=[2.2, 2.5, 4.0, 7.3],
)
para("Aprobaciones", bold=True, color=OSCURO)
grid_table(
    ["Rol", "Nombre", "Firma", "Fecha"],
    [
        ["Patrocinador", "", "", ""],
        ["Responsable funcional", "", "", ""],
        ["Líder técnico", "", "", ""],
    ],
    col_widths=[4.0, 5.0, 4.0, 3.0],
)

# ============================================================
# TABLA DE CONTENIDO (campo auto-actualizable)
# ============================================================
page_break()
h("Tabla de contenido", level=1)
para("Para actualizar: clic derecho sobre la tabla → «Actualizar campos» (o F9).",
     italic=True, size=9, color=GRIS)
p = doc.add_paragraph()
run = p.add_run()
fldStart = OxmlElement("w:fldChar"); fldStart.set(qn("w:fldCharType"), "begin")
instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
instr.text = 'TOC \\o "1-3" \\h \\z \\u'
fldSep = OxmlElement("w:fldChar"); fldSep.set(qn("w:fldCharType"), "separate")
t2 = OxmlElement("w:t"); t2.text = "Actualiza este campo para generar el índice."
fldEnd = OxmlElement("w:fldChar"); fldEnd.set(qn("w:fldCharType"), "end")
run._r.append(fldStart); run._r.append(instr); run._r.append(fldSep); run._r.append(t2); run._r.append(fldEnd)

# ============================================================
# 1. INTRODUCCIÓN
# ============================================================
page_break()
h("1. Introducción", level=1)

h("1.1 Propósito", level=2)
para("Este documento especifica los requisitos funcionales y no funcionales del "
     "Sistema de Gestión de Inventario de JJ Congeminco (Contratistas Generales "
     "en Minería y Construcción). Define qué debe hacer el sistema, las reglas de "
     "negocio que debe cumplir y las cualidades que debe satisfacer (seguridad, "
     "rendimiento, usabilidad, entre otras). Sirve como acuerdo entre los "
     "interesados y como base para diseño, construcción y pruebas.")

h("1.2 Alcance del producto", level=2)
para("El sistema centraliza el control de inventario de materiales, repuestos e "
     "insumos de la empresa, eliminando los registros dispersos en hojas de cálculo. "
     "Permite administrar un catálogo de productos, registrar movimientos (entradas, "
     "salidas, transferencias, existencias iniciales y ajustes) en múltiples "
     "almacenes, gestionar requerimientos de materiales con un flujo de aprobación, "
     "administrar órdenes de mantenimiento de vehículos y equipos con consumo de "
     "repuestos, generar reportes valorizados y de movimientos, e importar datos de "
     "forma masiva desde Excel.")
para("Fuera del alcance de esta línea base: facturación electrónica, integración "
     "contable externa, módulo de compras/órdenes de compra formales y aplicación "
     "móvil nativa (el sistema web es responsive).", italic=True, color=GRIS)

h("1.3 Definiciones, acrónimos y abreviaturas", level=2)
grid_table(
    ["Término", "Definición"],
    [
        ["SKU", "Código único interno que identifica a un producto."],
        ["Stock mínimo", "Umbral de existencias por debajo del cual se alerta para reabastecer."],
        ["Costo promedio", "Costo promedio ponderado móvil del producto (criterio NIC 2)."],
        ["Kardex", "Historial cronológico de movimientos de un producto con saldo corrido."],
        ["Ledger (libro mayor)", "Registro de movimientos append-only e inmutable; fuente de verdad del stock."],
        ["Saldo", "Existencia actual calculada por producto y ubicación; se almacena en caché."],
        ["Ubicación / Almacén", "Lugar físico donde se guarda el stock (almacén central, proyecto, otro)."],
        ["Requerimiento", "Solicitud formal de materiales asociada a un equipo o vehículo."],
        ["OT", "Orden de Trabajo de mantenimiento sobre un vehículo (placa)."],
        ["RBAC", "Control de acceso basado en roles."],
        ["RLS", "Row Level Security: seguridad a nivel de fila en la base de datos."],
        ["RF / RNF", "Requisito Funcional / Requisito No Funcional."],
        ["ERS / SRS", "Especificación de Requisitos de Software."],
        ["Compra directa", "Atención de un pedido sin stock: se registra entrada de compra y salida a la vez."],
    ],
    col_widths=[4.5, 11.5],
)

h("1.4 Referencias", level=2)
bullet("Código fuente del proyecto (monorepo): apps/web, packages/shared, packages/db.")
bullet("Contrato de tipos y validaciones: packages/shared/src/dto.ts y roles.ts.")
bullet("Modelo de datos y migraciones: packages/db (estándar BSG).")
bullet("Manual de Usuario del Sistema de Inventario JJ Congeminco, v1.0.")
bullet("ISO/IEC/IEEE 29148 e IEEE 830 como guía de estructura.")

h("1.5 Visión general del documento", level=2)
para("La sección 2 describe el contexto general del producto, su arquitectura y los "
     "tipos de usuario. La sección 3 detalla los requisitos funcionales agrupados por "
     "módulo. La sección 4 define los requisitos no funcionales. La sección 5 resume el "
     "modelo de datos. La sección 6 presenta la matriz de trazabilidad rol–permiso y los anexos.")

# ============================================================
# 2. DESCRIPCIÓN GENERAL
# ============================================================
page_break()
h("2. Descripción general", level=1)

h("2.1 Perspectiva del producto", level=2)
para("Es una aplicación web autónoma de tipo ERP de inventario, multi-almacén, "
     "accesible desde navegador en computador o celular. Se apoya en una base de "
     "datos relacional con un patrón de tres capas en el que el libro mayor de "
     "movimientos es la única fuente de verdad:")
para("Documento (cabecera + detalle)  →  Movimiento de stock (libro mayor inmutable)  →  Saldo (caché por producto y ubicación)",
     bold=True, color=OSCURO)

h("2.2 Arquitectura y tecnologías", level=2)
grid_table(
    ["Capa", "Tecnología", "Rol en el sistema"],
    [
        ["Frontend / App", "Next.js 15 (App Router), React 18, TypeScript", "Interfaz de usuario y rutas de la aplicación."],
        ["UI", "Tailwind CSS, shadcn/ui (Radix UI), lucide-react", "Componentes accesibles y diseño responsive."],
        ["Estado/Datos cliente", "TanStack Query (React Query)", "Caché de datos, sincronización y paginación."],
        ["Gráficos", "ECharts", "Visualizaciones del dashboard y reportes."],
        ["Validación", "Zod + React Hook Form", "Validación de formularios en cliente y servidor."],
        ["API", "Next.js Route Handlers (/app/api)", "Lógica de aplicación y orquestación con la BD."],
        ["Autenticación", "Supabase Auth (@supabase/ssr)", "Inicio de sesión y gestión de sesión."],
        ["Base de datos", "PostgreSQL (Supabase)", "Persistencia, funciones, triggers y RLS."],
        ["Archivos Excel", "SheetJS (xlsx)", "Lectura/generación de plantillas de importación."],
        ["Despliegue", "Cloudflare (OpenNext) + Wrangler", "Hospedaje y publicación del sitio."],
        ["Monorepo", "pnpm + Turborepo", "Gestión de paquetes y build (web, shared, db)."],
    ],
    col_widths=[3.2, 5.3, 7.5],
)

h("2.3 Funciones del producto (resumen)", level=2)
for f in [
    "Autenticación y control de acceso por rol y módulo (RBAC).",
    "Tablero (dashboard) con indicadores, gráficos y alertas de bajo mínimo.",
    "Consulta de saldos de stock con búsqueda y filtros.",
    "Administración del catálogo de productos con compatibilidad por tipo de equipo e imágenes.",
    "Registro de movimientos de inventario (5 tipos) con costeo y multi-almacén.",
    "Gestión de requerimientos de materiales con flujo de aprobación y separación de funciones.",
    "Mantenimiento de vehículos: órdenes de trabajo, consumo de repuestos y evidencia fotográfica.",
    "Reportes de movimientos, valorizado de inventario y recambios por desgaste.",
    "Importación masiva de productos y saldos desde Excel.",
    "Administración de maestros: categorías, personal, cargos, proveedores, almacenes, equipos, vehículos y tipos de equipo.",
]:
    bullet(f)

h("2.4 Características de los usuarios (roles)", level=2)
grid_table(
    ["Rol", "Perfil", "Responsabilidad principal"],
    [
        ["Administrador (admin)", "Administrador del sistema", "Configuración total, importación, maestros y operación."],
        ["Gerencia (gerencia)", "Jefatura", "Consulta, aprobación de requerimientos y reportes."],
        ["Supervisión (supervision)", "Supervisor de campo", "Registra movimientos; crea y aprueba requerimientos."],
        ["Almacenero (almacenero)", "Encargado de almacén", "Operación diaria: catálogo, movimientos, mantenimiento."],
        ["Logística (logistica)", "Consulta", "Solo visualización de stock, movimientos y reportes."],
    ],
    col_widths=[4.2, 4.0, 7.8],
)
para("La asignación rol→módulo reside en la base de datos (configurable sin "
     "redesplegar). El frontend recibe los módulos permitidos y filtra el menú; la "
     "API y la RLS refuerzan el control.", italic=True, color=GRIS)

h("2.5 Restricciones de diseño", level=2)
bullet("La base de datos es la fuente de verdad: las reglas críticas se aplican con funciones, triggers y RLS de PostgreSQL.", "Defensa en BD: ")
bullet("El libro mayor de movimientos es inmutable; no se edita ni borra. Para corregir se anula el documento.", "Inmutabilidad: ")
bullet("Nomenclatura y estructura según el estándar BSG (objetos en PascalCase español, esquemas comun/seg/inv).", "Estándar BSG: ")
bullet("Toda PK es UUID por compatibilidad con Supabase Auth.", "Identificadores: ")
bullet("Importaciones limitadas a 5000 filas por archivo y validación todo-o-nada.", "Importación: ")

h("2.6 Suposiciones y dependencias", level=2)
bullet("Disponibilidad del servicio Supabase (Auth + PostgreSQL + Storage).")
bullet("Conectividad a Internet por parte de los usuarios.")
bullet("Los esquemas inv y seg deben estar expuestos a la Data API de Supabase.")
bullet("Cada usuario de Auth debe tener su fila correspondiente en seg.T_Usuario para poder operar.")
bullet("Montos expresados en Soles (S/), formato peruano.")

# ============================================================
# 3. REQUISITOS FUNCIONALES
# ============================================================
page_break()
h("3. Requisitos funcionales", level=1)
para("Cada requisito se identifica con un código del tipo RF-XXX-NN. La prioridad "
     "sigue la convención MoSCoW (Imprescindible / Recomendable / Opcional).",
     italic=True, color=GRIS)

# ---- 3.1 Autenticación ----
h("3.1 Autenticación y control de acceso", level=2)
rf_block("RF-AUT-01", "Inicio de sesión", "Imprescindible",
    "El sistema debe permitir iniciar sesión con correo electrónico y contraseña mediante Supabase Auth.",
    "Usuario no autenticado",
    "El usuario debe existir en el proveedor de autenticación.",
    ["Credenciales inválidas muestran mensaje de error sin revelar detalles.",
     "Tras autenticarse, se redirige al tablero principal."])
rf_block("RF-AUT-02", "Cierre de sesión", "Imprescindible",
    "El sistema debe permitir cerrar la sesión desde el menú de usuario.",
    "Usuario autenticado", "Sesión activa.",
    ["Al cerrar sesión se invalida la sesión y se redirige a /login."])
rf_block("RF-AUT-03", "Gestión y refresco de sesión", "Imprescindible",
    "El sistema debe mantener y refrescar la sesión del usuario de forma transparente mediante middleware.",
    "Usuario autenticado", "Sesión válida.",
    ["Si la sesión expira o no es válida, se redirige a /login."])
rf_block("RF-AUT-04", "Control de acceso por rol y módulo (RBAC)", "Imprescindible",
    "El sistema debe mostrar y permitir solo los módulos habilitados para el rol del usuario.",
    "Todos los roles", "Usuario con rol asignado.",
    ["El menú lateral se filtra según los módulos del usuario.",
     "El acceso por URL a un módulo no permitido redirige al tablero.",
     "La restricción se refuerza en la API y en la RLS de la BD."])
rf_block("RF-AUT-05", "Vinculación de usuario operativo", "Imprescindible",
    "Un usuario de autenticación no puede operar hasta tener su registro en seg.T_Usuario con un rol.",
    "Administrador", "Usuario creado en Auth.",
    ["Sin fila en seg.T_Usuario, el acceso a datos falla por RLS."])

# ---- 3.2 Dashboard ----
h("3.2 Tablero (Dashboard)", level=2)
rf_block("RF-DSH-01", "Indicadores clave", "Imprescindible",
    "Mostrar total de productos activos, valor total del inventario, cantidad bajo mínimo y movimientos del período.",
    "Todos los roles con módulo dashboard", "Sesión iniciada.",
    ["El valor total se calcula como stock × costo promedio.",
     "Solo lectura."])
rf_block("RF-DSH-02", "Gráficos analíticos", "Recomendable",
    "Mostrar gráficos de entradas vs salidas por día, valor por categoría y top de productos más movidos.",
    "Todos los roles con módulo dashboard", "Existencia de movimientos.",
    ["Los gráficos respetan el período seleccionado."])
rf_block("RF-DSH-03", "Tabla de productos bajo mínimo", "Imprescindible",
    "Listar los productos cuyo stock total es menor a su stock mínimo (SKU, nombre, categoría, mínimo, actual).",
    "Todos los roles con módulo dashboard", "Catálogo con stock mínimo configurado.",
    ["Marca visual de estado «Bajo mínimo»."])
rf_block("RF-DSH-04", "Selector de período", "Recomendable",
    "Permitir filtrar la información del tablero por últimos 7, 30 o 90 días.",
    "Todos los roles con módulo dashboard", "—",
    ["Todos los indicadores y gráficos se recalculan con el rango."])

# ---- 3.3 Saldos ----
h("3.3 Saldos", level=2)
rf_block("RF-SLD-01", "Consulta de saldos", "Imprescindible",
    "Mostrar el stock disponible de cada producto en tarjetas, optimizado para uso en celular/campo.",
    "Todos los roles con módulo saldos", "Productos en catálogo.",
    ["Muestra cantidad con su unidad de medida.",
     "Marca «Bajo mínimo» cuando corresponde."])
rf_block("RF-SLD-02", "Búsqueda", "Imprescindible",
    "Permitir buscar productos por nombre o SKU.",
    "Todos los roles con módulo saldos", "—",
    ["Búsqueda incremental."])
rf_block("RF-SLD-03", "Filtro por categoría", "Recomendable",
    "Permitir filtrar saldos por categoría mediante selección rápida.",
    "Todos los roles con módulo saldos", "Categorías definidas.", ["—"])
rf_block("RF-SLD-04", "Detalle de producto", "Recomendable",
    "Al abrir un producto, mostrar stock total y por ubicación, stock mínimo, costo promedio y equipos compatibles.",
    "Todos los roles con módulo saldos", "—",
    ["Solo consulta; no permite modificar stock."])

# ---- 3.4 Catálogo ----
h("3.4 Catálogo de productos", level=2)
rf_block("RF-CAT-01", "Listar, buscar y filtrar productos", "Imprescindible",
    "Listar el catálogo con SKU, nombre, categoría, tipos de equipo compatibles y estado de stock; con búsqueda y filtro por categoría.",
    "Todos los roles con módulo catálogo", "—", ["—"])
rf_block("RF-CAT-02", "Crear producto", "Imprescindible",
    "Registrar un producto con SKU, nombre, categoría, unidad de medida, stock mínimo, códigos y compatibilidad.",
    "admin, almacenero", "Categorías y unidades existentes.",
    ["SKU obligatorio (máx 50) y único; nombre obligatorio (máx 200).",
     "Categoría y unidad obligatorias; stock mínimo ≥ 0.",
     "Compatibilidad: «General» XOR uno o más tipos de equipo (nunca ambos)."])
rf_block("RF-CAT-03", "Editar producto", "Imprescindible",
    "Modificar los datos de un producto existente, incluida su compatibilidad y estado.",
    "admin, almacenero", "Producto existente.",
    ["Se mantienen las mismas validaciones de creación."])
rf_block("RF-CAT-04", "Eliminar producto", "Recomendable",
    "Eliminar un producto del catálogo.",
    "admin, almacenero", "Producto sin dependencias.",
    ["No se permite eliminar si tiene movimientos u otras dependencias."])
rf_block("RF-CAT-05", "Gestionar imágenes", "Opcional",
    "Adjuntar hasta 3 imágenes por producto, marcando una como principal.",
    "admin, almacenero", "Producto existente.",
    ["Máximo 3 imágenes; orden y marca de principal."])
rf_block("RF-CAT-06", "Compatibilidad por tipo de equipo", "Imprescindible",
    "Definir si un producto es general (todos los equipos) o aplica a tipos de equipo específicos.",
    "admin, almacenero", "Tipos de equipo definidos.",
    ["Invariante general XOR ≥1 tipo validada por la BD.",
     "Producto no general sin tipo aparece como «Sin clasificar»."])
rf_block("RF-CAT-07", "Asociación masiva por categoría", "Opcional",
    "Asociar de una sola vez todos los productos de una categoría a un tipo de equipo.",
    "admin, almacenero", "Categoría y tipo existentes.", ["—"])
rf_block("RF-CAT-08", "Kardex del producto", "Recomendable",
    "Consultar el historial de movimientos de un producto con saldo corrido por ubicación.",
    "Todos los roles con módulo catálogo", "Movimientos registrados.", ["Solo lectura."])

# ---- 3.5 Movimientos ----
h("3.5 Movimientos de inventario", level=2)
rf_block("RF-MOV-01", "Registrar documento de inventario", "Imprescindible",
    "Registrar un documento con cabecera y una o más líneas de detalle que afecta el stock al confirmarse.",
    "admin, almacenero, supervisión", "Productos y ubicaciones definidos.",
    ["El detalle requiere al menos una línea con cantidad > 0.",
     "Al registrar se confirma e impacta el libro mayor de inmediato."])
rf_block("RF-MOV-02", "Tipos de documento", "Imprescindible",
    "Soportar los tipos: existencia inicial, entrada, salida, transferencia y ajuste.",
    "admin, almacenero, supervisión", "—",
    ["Entrada/existencia inicial: requieren ubicación destino.",
     "Salida: requiere origen y placa por línea.",
     "Transferencia: requiere origen y destino distintos.",
     "Ajuste: corrige por conteo físico."])
rf_block("RF-MOV-03", "Validaciones de negocio", "Imprescindible",
    "Aplicar reglas que garanticen la integridad del stock.",
    "admin, almacenero, supervisión", "—",
    ["Ninguna salida puede dejar saldo negativo (se bloquea).",
     "Transferencia con origen ≠ destino.",
     "Cada línea de salida exige placa destino."])
rf_block("RF-MOV-04", "Costeo de movimientos", "Imprescindible",
    "Gestionar el costo unitario según el tipo de movimiento.",
    "admin, almacenero, supervisión", "—",
    ["Entrada: costo unitario editable.",
     "Salida: costo = costo promedio vigente (override manual registrado).",
     "Transferencia: mantiene el costo; no recalcula el promedio."])
rf_block("RF-MOV-05", "Filtro de compatibilidad por placa", "Opcional",
    "En salidas, permitir mostrar solo productos compatibles con la placa (más los generales).",
    "admin, almacenero, supervisión", "Vehículo con tipo de equipo.", ["—"])
rf_block("RF-MOV-06", "Documentos recientes", "Recomendable",
    "Listar los documentos registrados recientemente, con paginación.",
    "admin, almacenero, supervisión", "—", ["—"])

# ---- 3.6 Requerimientos ----
h("3.6 Requerimientos", level=2)
rf_block("RF-REQ-01", "Crear requerimiento", "Imprescindible",
    "Registrar una solicitud de materiales asociada a un equipo o vehículo, con una o más líneas.",
    "admin, almacenero, supervisión", "Productos y destino definidos.",
    ["Debe indicarse equipo o placa (cabecera) o placa por línea.",
     "Al crearse queda en estado «pendiente» y no descuenta stock."])
rf_block("RF-REQ-02", "Origen del requerimiento", "Recomendable",
    "Clasificar el requerimiento como planificado, presupuestado o por desgaste prematuro.",
    "admin, almacenero, supervisión", "—",
    ["El origen «desgaste prematuro» alimenta el reporte de recambios."])

# ---- 3.7 Aprobaciones ----
h("3.7 Aprobaciones", level=2)
rf_block("RF-APR-01", "Bandeja de pendientes", "Imprescindible",
    "Listar los requerimientos pendientes de atención o rechazo.",
    "admin, gerencia, supervisión", "Requerimientos pendientes.", ["—"])
rf_block("RF-APR-02", "Atender requerimiento", "Imprescindible",
    "Entregar el material indicando almacén origen y cantidad por línea (entrega parcial permitida).",
    "admin, gerencia, supervisión", "Requerimiento pendiente.",
    ["Modo por línea: stock (sale del almacén) o compra directa.",
     "Compra directa exige proveedor, comprobante y costo por línea.",
     "Al atender genera salida valorizada y marca «atendido»; actualiza cantidad atendida."])
rf_block("RF-APR-03", "Rechazar requerimiento", "Imprescindible",
    "Rechazar un requerimiento pendiente con motivo opcional.",
    "admin, gerencia, supervisión", "Requerimiento pendiente.",
    ["El requerimiento pasa a estado «anulado»; no descuenta stock."])
rf_block("RF-APR-04", "Separación de funciones", "Imprescindible",
    "Impedir que el creador de un requerimiento lo apruebe él mismo.",
    "Todos los aprobadores", "—",
    ["El administrador está exento. La regla la refuerza la BD."])
rf_block("RF-APR-05", "Generar PDF de la solicitud", "Recomendable",
    "Generar un PDF del requerimiento para impresión o archivo.",
    "admin, gerencia, supervisión", "Requerimiento existente.", ["—"])
rf_block("RF-APR-06", "Reconciliación de mantenimiento", "Imprescindible",
    "Ratificar (aprobar) o rechazar el consumo de repuestos de órdenes de trabajo.",
    "admin, gerencia, supervisión", "OT en estado «consumida».",
    ["Aprobar: cierra la OT (el stock ya salió).",
     "Rechazar: anula la OT y genera entrada de reversa al costo del libro mayor."])
rf_block("RF-APR-07", "Histórico", "Recomendable",
    "Consultar requerimientos atendidos y anulados (solo lectura).",
    "admin, gerencia, supervisión", "—", ["—"])

# ---- 3.8 Mantenimiento ----
h("3.8 Mantenimiento (Órdenes de Trabajo)", level=2)
rf_block("RF-MNT-01", "Crear orden de trabajo", "Imprescindible",
    "Registrar una OT por placa con tipo, fecha, turno, mecánico(s), kilometraje y trabajos.",
    "admin, almacenero, supervisión", "Vehículos y personal definidos.",
    ["Tipo preventivo o correctivo; placa y al menos un personal obligatorios."])
rf_block("RF-MNT-02", "Ciclo de estados", "Imprescindible",
    "Gestionar los estados de la OT: abierta, consumida, cerrada y anulada.",
    "admin, almacenero, supervisión", "OT existente.",
    ["abierta → consumida → cerrada | anulada."])
rf_block("RF-MNT-03", "Consumir repuestos", "Imprescindible",
    "Registrar el consumo de repuestos, descontando el stock de inmediato.",
    "admin, almacenero, supervisión", "OT abierta y stock disponible.",
    ["Modo stock o compra directa (proveedor, comprobante y costo por línea).",
     "Genera un requerimiento interno enlazado 1:1 a la OT.",
     "La OT pasa a «consumida», pendiente de reconciliación."])
rf_block("RF-MNT-04", "Cerrar sin repuestos", "Recomendable",
    "Cerrar una OT que no consumió repuestos.",
    "admin, almacenero, supervisión", "OT abierta.", ["La OT pasa a «cerrada»."])
rf_block("RF-MNT-05", "Anular / eliminar OT", "Recomendable",
    "Anular una OT o eliminarla si está abierta y sin consumo.",
    "admin, almacenero, supervisión", "OT abierta sin consumo (para eliminar).", ["—"])
rf_block("RF-MNT-06", "Evidencia fotográfica", "Opcional",
    "Adjuntar fotos de estado actual (antes) y post mantenimiento (después) al culminar.",
    "admin, almacenero, supervisión", "OT existente.",
    ["Mínimo 1 de cada tipo al cerrar; máximo 10 por tipo."])
rf_block("RF-MNT-07", "Generar PDF de la OT", "Recomendable",
    "Imprimir/generar el PDF de la orden de trabajo.",
    "admin, almacenero, supervisión", "OT existente.", ["—"])

# ---- 3.9 Reportes ----
h("3.9 Reportes", level=2)
rf_block("RF-REP-01", "Reporte de movimientos", "Imprescindible",
    "Analizar entradas y salidas por período con filtros y totales por tipo.",
    "admin, gerencia, supervisión", "Movimientos registrados.",
    ["Filtros por fecha, producto, categoría, proveedor, ubicación, placa, equipo y tipo.",
     "Indicadores de valor de entradas, salidas y neto.",
     "Exportable a CSV."])
rf_block("RF-REP-02", "Reporte valorizado", "Imprescindible",
    "Mostrar el valor del inventario (stock × costo promedio) por categoría.",
    "admin, gerencia, supervisión", "Catálogo con stock.",
    ["Filtro por categoría y por «solo bajo mínimo».",
     "Exportable a CSV."])
rf_block("RF-REP-03", "Reporte de recambios", "Recomendable",
    "Detectar repuestos cambiados con mayor frecuencia de lo normal (desgaste prematuro).",
    "admin, gerencia, supervisión", "Requerimientos/consumos con origen de desgaste.",
    ["Indicadores de total de casos, acelerados y porcentaje."])
rf_block("RF-REP-04", "Exportación CSV", "Recomendable",
    "Permitir descargar los reportes en formato CSV.",
    "admin, gerencia, supervisión", "Reporte generado.", ["—"])

# ---- 3.10 Importación ----
h("3.10 Importación masiva", level=2)
rf_block("RF-IMP-01", "Importar productos", "Imprescindible",
    "Cargar productos de forma masiva desde un archivo Excel (.xlsx).",
    "admin", "Maestros de categorías, unidades y tipos definidos.",
    ["Columnas: Sku, Nombre, CodigoCategoria, CodigoUnidad, EsGeneral, TiposEquipo, StockMinimo, CodigoBarra, CodigoProductoProveedor.",
     "Modo «solo crear» o «crear y actualizar»."])
rf_block("RF-IMP-02", "Importar saldos", "Imprescindible",
    "Cargar saldos por almacén desde Excel, en modo existencia inicial o recuento.",
    "admin", "Productos y ubicaciones existentes; fecha de corte.",
    ["Inicial: crea existencias (rechaza si ya hay saldo).",
     "Recuento: ajusta la diferencia contra el saldo vigente.",
     "Requiere fecha de corte (YYYY-MM-DD)."])
rf_block("RF-IMP-03", "Validación todo-o-nada", "Imprescindible",
    "Si una fila falla, no se aplica ningún cambio y se reporta el detalle por fila.",
    "admin", "Archivo cargado.",
    ["Reporte con filas correctas, creados, actualizados y errores (fila/columna)."])
rf_block("RF-IMP-04", "Plantilla descargable", "Recomendable",
    "Proveer una plantilla Excel con el formato correcto y los códigos de referencia.",
    "admin", "—", ["—"])
rf_block("RF-IMP-05", "Límite de tamaño", "Imprescindible",
    "Limitar la importación a un máximo de 5000 filas por archivo.",
    "admin", "—", ["Se rechazan archivos que excedan el límite."])

# ---- 3.11 Maestros ----
h("3.11 Maestros (configuración)", level=2)
para("Cada maestro debe permitir listar, crear, editar y eliminar registros "
     "(eliminación condicionada a que no existan dependencias). El acceso de "
     "escritura sigue la matriz de permisos del Anexo 6.2.")
grid_table(
    ["Código", "Maestro", "Campos principales", "Escritura"],
    [
        ["RF-MAE-01", "Categorías y familias", "Código, nombre, descripción, familia padre", "admin"],
        ["RF-MAE-02", "Personal", "Nombre, DNI, teléfono, cargo, usuario de acceso", "admin"],
        ["RF-MAE-03", "Cargos", "Código, nombre, descripción", "admin"],
        ["RF-MAE-04", "Proveedores", "Nombre, RUC, contacto, teléfono, cuentas bancarias", "admin, almacenero"],
        ["RF-MAE-05", "Almacenes / ubicaciones", "Código, nombre, tipo, dirección", "admin"],
        ["RF-MAE-06", "Equipos", "Código, nombre, descripción, tipo de equipo", "admin, almacenero"],
        ["RF-MAE-07", "Vehículos", "Placa, modelo, equipo asignado", "admin, almacenero"],
        ["RF-MAE-08", "Tipos de equipo", "Código, nombre, descripción", "admin, almacenero"],
    ],
    col_widths=[2.4, 4.2, 6.6, 2.8],
)
para("Detalle del maestro de proveedores: cada proveedor admite varias cuentas "
     "bancarias (banco, tipo corriente/ahorros, número, CCI, moneda PEN/USD, titular, "
     "marca de principal).", italic=True, color=GRIS)

# ============================================================
# 4. REQUISITOS NO FUNCIONALES
# ============================================================
page_break()
h("4. Requisitos no funcionales", level=1)

def rnf(rid, nombre, texto, prioridad="Imprescindible"):
    p = doc.add_paragraph()
    r = p.add_run(f"{rid} — {nombre} ")
    r.bold = True
    r.font.color.rgb = OSCURO
    r.font.size = Pt(10.5)
    r2 = p.add_run(f"[{prioridad}]")
    r2.italic = True
    r2.font.size = Pt(9)
    r2.font.color.rgb = NARANJA
    para(texto, space_after=8)

h("4.1 Seguridad", level=2)
rnf("RNF-SEG-01", "Autenticación", "El acceso debe requerir autenticación mediante Supabase Auth; ninguna ruta de la aplicación es accesible sin sesión válida.")
rnf("RNF-SEG-02", "Seguridad a nivel de fila (RLS)", "La base de datos debe proteger los datos con políticas RLS por rol, de modo que cada usuario solo pueda leer/escribir lo que su rol permite, independientemente del frontend.")
rnf("RNF-SEG-03", "Defensa en profundidad", "El control de acceso debe aplicarse en tres capas: menú/guard de ruta en el frontend, verificación en la API y RLS en la base de datos.")
rnf("RNF-SEG-04", "Separación de funciones", "El sistema debe impedir que un mismo usuario cree y apruebe el mismo requerimiento (salvo el administrador).")
rnf("RNF-SEG-05", "Validación de entrada", "Toda entrada debe validarse con Zod tanto en el cliente como en el servidor antes de persistir.")
rnf("RNF-SEG-06", "Transporte cifrado", "Toda comunicación debe realizarse sobre HTTPS.")
rnf("RNF-SEG-07", "Gestión de secretos", "Las credenciales y llaves (Supabase, base de datos) deben gestionarse por variables de entorno y nunca exponerse en el cliente (solo la llave anónima).")

h("4.2 Integridad y fiabilidad de datos", level=2)
rnf("RNF-FIA-01", "Fuente única de verdad", "El stock debe derivarse siempre del libro mayor de movimientos (append-only e inmutable); el saldo es una caché reconciliable.")
rnf("RNF-FIA-02", "Inmutabilidad del libro mayor", "Los movimientos no deben poder editarse ni borrarse; las correcciones se hacen por anulación de documento (movimientos inversos).")
rnf("RNF-FIA-03", "Operaciones atómicas", "El registro de documentos y las importaciones deben ser transaccionales (todo-o-nada).")
rnf("RNF-FIA-04", "No negatividad del stock", "El sistema debe impedir cualquier movimiento que deje un saldo negativo por producto y ubicación.")
rnf("RNF-FIA-05", "Reconciliación", "Debe existir un mecanismo para verificar la caché de saldos contra el libro mayor.")
rnf("RNF-FIA-06", "Auditoría", "Cada registro debe conservar usuario y fecha de creación/modificación, versión de fila y trazabilidad (campos de auditoría del estándar BSG).")

h("4.3 Rendimiento", level=2)
rnf("RNF-REN-01", "Consulta de saldos eficiente", "La consulta de stock debe resolverse en tiempo prácticamente constante gracias a la caché de saldos por producto y ubicación.", "Imprescindible")
rnf("RNF-REN-02", "Caché de cliente", "El frontend debe cachear y reutilizar datos con TanStack Query para minimizar llamadas y mejorar la respuesta percibida.", "Recomendable")
rnf("RNF-REN-03", "Paginación", "Los listados extensos (documentos, maestros) deben paginarse para mantener tiempos de carga aceptables.", "Recomendable")
rnf("RNF-REN-04", "Límite de lotes", "Las importaciones se limitan a 5000 filas por archivo para acotar el consumo de recursos.", "Imprescindible")

h("4.4 Usabilidad", level=2)
rnf("RNF-USA-01", "Diseño responsive", "La interfaz debe adaptarse a computador y celular; la pantalla de saldos debe estar optimizada para uso en campo.")
rnf("RNF-USA-02", "Idioma", "Toda la interfaz y los mensajes deben estar en español.")
rnf("RNF-USA-03", "Tema claro/oscuro", "El usuario debe poder alternar entre tema claro y oscuro.", "Opcional")
rnf("RNF-USA-04", "Retroalimentación", "El sistema debe informar el resultado de cada acción mediante notificaciones claras (éxito/error).")
rnf("RNF-USA-05", "Accesibilidad", "Los componentes de interfaz deben basarse en primitivas accesibles (Radix UI).", "Recomendable")
rnf("RNF-USA-06", "Mensajes de error útiles", "Los errores de validación deben explicar la causa y, cuando aplique, la fila/columna afectada.")

h("4.5 Mantenibilidad y calidad", level=2)
rnf("RNF-MAN-01", "Tipado estricto", "El código debe usar TypeScript con tipos compartidos entre frontend y backend (paquete shared).")
rnf("RNF-MAN-02", "Contrato único", "Las reglas de validación y los enumerados deben centralizarse en un único contrato (Zod) reutilizado por todas las capas.")
rnf("RNF-MAN-03", "Estándar de base de datos", "El esquema debe seguir el estándar BSG (nomenclatura, esquemas, auditoría, constraints nombradas).")
rnf("RNF-MAN-04", "Migraciones idempotentes", "Los cambios de base de datos deben aplicarse mediante migraciones versionadas e idempotentes.")
rnf("RNF-MAN-05", "Monorepo", "El proyecto debe organizarse como monorepo (pnpm + Turborepo) separando web, shared y db.")

h("4.6 Portabilidad y despliegue", level=2)
rnf("RNF-POR-01", "Despliegue web", "El sistema debe poder desplegarse en Cloudflare mediante OpenNext/Wrangler.")
rnf("RNF-POR-02", "Compatibilidad de navegadores", "Debe funcionar en navegadores modernos (Chrome, Edge, Safari, Firefox) en sus versiones recientes.")
rnf("RNF-POR-03", "Multi-almacén", "La arquitectura debe soportar múltiples almacenes/ubicaciones desde el inicio.")
rnf("RNF-POR-04", "Independencia de proveedor de datos", "El acceso a datos debe concentrarse en una capa que facilite cambios futuros de infraestructura.", "Opcional")

h("4.7 Cumplimiento y reglas contables", level=2)
rnf("RNF-LEG-01", "Costeo NIC 2", "La valorización de existencias debe usar costo promedio ponderado (criterio NIC 2).")
rnf("RNF-LEG-02", "Moneda local", "Los montos deben expresarse en Soles (S/) con formato peruano.")
rnf("RNF-LEG-03", "Soporte de comprobantes", "Los movimientos deben permitir registrar número de comprobante (factura, boleta, guía).", "Recomendable")

h("4.8 Disponibilidad y respaldo", level=2)
rnf("RNF-DIS-01", "Disponibilidad", "El sistema debe estar disponible en horario operativo de la empresa, sujeto a la disponibilidad del proveedor de nube.", "Recomendable")
rnf("RNF-DIS-02", "Respaldo de datos", "La base de datos debe contar con respaldos automáticos provistos por la plataforma gestionada.", "Recomendable")

# ============================================================
# 5. MODELO DE DATOS (RESUMEN)
# ============================================================
page_break()
h("5. Modelo de datos (resumen)", level=1)
para("El modelo sigue el patrón ERP de tres capas. A continuación, las entidades "
     "principales agrupadas por esquema.")
grid_table(
    ["Esquema", "Entidad", "Descripción"],
    [
        ["seg", "T_Rol", "Roles del sistema."],
        ["seg", "T_Usuario", "Usuarios operativos vinculados a Auth, con su rol."],
        ["seg", "T_RolModulo", "Asignación de módulos visibles por rol (RBAC configurable)."],
        ["comun", "T_UnidadMedida", "Unidades de medida (LT, KG, UND…)."],
        ["inv", "T_Categoria", "Categorías/familias jerárquicas de productos."],
        ["inv", "T_Producto", "Catálogo de productos (con compatibilidad e imágenes)."],
        ["inv", "T_TipoEquipo", "Tipos de equipo para compatibilidad de repuestos."],
        ["inv", "T_Equipo / T_Vehiculo", "Equipos y vehículos (placas)."],
        ["inv", "T_Ubicacion", "Almacenes/ubicaciones (multi-almacén)."],
        ["inv", "T_Proveedor", "Proveedores y sus cuentas bancarias."],
        ["inv", "T_DocumentoInventario", "Cabecera de documentos de inventario."],
        ["inv", "T_DocumentoInventarioDetalle", "Líneas de cada documento."],
        ["inv", "T_MovimientoStock", "Libro mayor de movimientos (append-only, inmutable)."],
        ["inv", "T_SaldoStock", "Caché de saldos por producto y ubicación."],
        ["inv", "T_Requerimiento / Detalle", "Requerimientos de materiales y sus líneas."],
        ["inv", "T_OrdenMantenimiento", "Órdenes de trabajo y consumos asociados."],
        ["inv", "T_Importacion", "Registro de importaciones masivas."],
    ],
    col_widths=[2.2, 5.4, 8.4],
)
para("Funciones clave de la BD: FnConfirmarDocumentoInventario (genera movimientos), "
     "FnGuardarProducto (valida compatibilidad), FnImportarProductos / "
     "FnImportarSaldosIniciales (importación validada), FnModulosUsuario y "
     "FnRolUsuario (control de acceso). Triggers mantienen la caché de saldos, "
     "el costo promedio y la auditoría.", italic=True, color=GRIS)

# ============================================================
# 6. ANEXOS
# ============================================================
page_break()
h("6. Anexos", level=1)

h("6.1 Matriz de acceso por módulo (rol → módulo)", level=2)
grid_table(
    ["Módulo", "Admin", "Gerencia", "Supervisión", "Almacenero"],
    [
        ["Dashboard", "Sí", "Sí", "Sí", "Sí"],
        ["Saldos", "Sí", "Sí", "Sí", "Sí"],
        ["Catálogo", "Sí", "Sí", "Sí", "Sí"],
        ["Movimientos", "Sí", "Sí", "Sí", "Sí"],
        ["Requerimientos", "Sí", "Sí", "Sí", "Sí"],
        ["Mantenimiento", "Sí", "Sí", "Sí", "Sí"],
        ["Aprobaciones", "Sí", "Sí", "Sí", "No"],
        ["Reportes", "Sí", "Sí", "Sí", "No"],
        ["Importar", "Sí", "No", "No", "No"],
        ["Maestros (general)", "Sí", "Sí", "Sí", "No"],
        ["Maestros · Proveedores", "Sí", "Sí", "No", "No"],
    ],
    col_widths=[5.0, 2.7, 2.7, 2.9, 2.7],
)

h("6.2 Matriz de permisos de escritura (rol → recurso)", level=2)
grid_table(
    ["Permiso", "Recurso protegido", "Roles con escritura"],
    [
        ["productoEscritura", "Productos, imágenes, equipos, vehículos, tipos, proveedores", "admin, almacenero"],
        ["documentoEscritura", "Documentos de inventario (entradas/salidas/transferencias)", "admin, almacenero, supervisión"],
        ["catalogoAdmin", "Categorías, almacenes, cargos, personal, usuarios, importación", "admin"],
        ["requerimientoCrear", "Crear requerimientos y órdenes de mantenimiento", "admin, almacenero, supervisión"],
        ["requerimientoAprobar", "Atender/rechazar requerimientos; reconciliar OT", "admin, gerencia, supervisión"],
    ],
    col_widths=[3.6, 7.2, 5.2],
)

h("6.3 Diccionario de estados", level=2)
grid_table(
    ["Entidad", "Estados", "Transiciones"],
    [
        ["Documento de inventario", "borrador → confirmado", "Se confirma al registrar; se revierte por anulación."],
        ["Requerimiento", "pendiente → atendido | anulado", "Atender (con entrega) o rechazar."],
        ["Orden de mantenimiento", "abierta → consumida → cerrada | anulada", "Consumir, reconciliar (cerrar) o rechazar (anular + reversa)."],
    ],
    col_widths=[4.2, 5.4, 6.4],
)

h("6.4 Reglas de negocio destacadas", level=2)
for r in [
    "El libro mayor de movimientos es inmutable; toda corrección se hace por anulación.",
    "Ningún movimiento puede dejar saldo negativo por producto y ubicación.",
    "Un producto es «general» (todos los equipos) o aplica a uno o más tipos, nunca ambos.",
    "Las salidas exigen placa destino; las transferencias, origen y destino distintos.",
    "Las salidas se valorizan al costo promedio vigente; la compra directa recalcula el promedio resultante.",
    "Quien crea un requerimiento no puede aprobarlo (salvo el administrador).",
    "El consumo de repuestos descuenta stock de inmediato y queda pendiente de reconciliación.",
    "Las importaciones son todo-o-nada con reporte de errores por fila.",
]:
    bullet(r)

# Pie final
doc.add_paragraph()
para("— Fin del documento —", italic=True, color=GRIS, align=WD_ALIGN_PARAGRAPH.CENTER)
para("Especificación de Requisitos · Sistema de Gestión de Inventario JJ Congeminco · v1.0 · Junio 2026",
     size=8, color=GRIS, align=WD_ALIGN_PARAGRAPH.CENTER)

# ---------- Pie de página con numeración ----------
section = doc.sections[0]
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run("JJ Congeminco — ERS Sistema de Inventario   |   Página ")
run.font.size = Pt(8); run.font.color.rgb = GRIS
# campo PAGE
fld1 = OxmlElement("w:fldSimple"); fld1.set(qn("w:instr"), "PAGE")
fp._p.append(fld1)

import os
out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                   "ERS-Sistema-Inventario-Congeminco.docx")
doc.save(out)
print("OK:", out)
